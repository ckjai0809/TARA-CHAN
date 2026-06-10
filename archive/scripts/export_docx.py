"""Convert PRESENTATION_SUMMARY.md to a formatted Word document."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


ROOT = Path(__file__).parent.parent
SRC = ROOT / "PRESENTATION_SUMMARY.md"
OUT = ROOT / "PRESENTATION_SUMMARY.docx"


def set_heading_style(paragraph, level):
    colors = {
        1: RGBColor(0x1A, 0x1A, 0x2E),
        2: RGBColor(0x16, 0x21, 0x3E),
        3: RGBColor(0x0F, 0x3C, 0x78),
    }
    sizes = {1: 24, 2: 18, 3: 14}
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run("")
    run.font.size = Pt(sizes.get(level, 12))
    run.font.color.rgb = colors.get(level, RGBColor(0, 0, 0))
    run.font.bold = True
    if level == 1:
        paragraph.paragraph_format.space_before = Pt(18)
        paragraph.paragraph_format.space_after = Pt(6)
    elif level == 2:
        paragraph.paragraph_format.space_before = Pt(14)
        paragraph.paragraph_format.space_after = Pt(4)
    else:
        paragraph.paragraph_format.space_before = Pt(10)
        paragraph.paragraph_format.space_after = Pt(2)


def add_table_from_lines(doc, header_line, separator_line, data_lines):
    all_lines = [header_line] + data_lines
    rows = []
    for line in all_lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)

    if not rows:
        return

    ncols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            # strip inline bold markers for table cells
            cell_text = re.sub(r"\*\*(.+?)\*\*", r"\1", cell_text)
            cell.text = cell_text
            if i == 0:
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                # shade header row
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "1F3C78")
                tcPr.append(shd)

    doc.add_paragraph()  # spacing after table


def inline_bold(paragraph, text):
    """Add a paragraph run with **bold** markers rendered as bold."""
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def render_code_block(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("\n".join(lines))
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)


def main():
    raw = SRC.read_text(encoding="utf-8")
    lines = raw.splitlines()

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Default body font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    i = 0
    in_code_block = False
    code_lines = []

    # collect table state
    pending_header = None
    pending_sep = None

    while i < len(lines):
        line = lines[i]

        # Code fence
        if line.strip().startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                in_code_block = False
                render_code_block(doc, code_lines)
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Table detection
        if line.strip().startswith("|"):
            # check if next line is separator
            if i + 1 < len(lines) and re.match(r"^\|[-| :]+\|", lines[i + 1].strip()):
                pending_header = line
                pending_sep = lines[i + 1]
                data = []
                j = i + 2
                while j < len(lines) and lines[j].strip().startswith("|"):
                    data.append(lines[j])
                    j += 1
                add_table_from_lines(doc, pending_header, pending_sep, data)
                i = j
                pending_header = pending_sep = None
                continue
            else:
                # continuation row without separator — treat as text
                pass

        # Headings
        if line.startswith("### "):
            p = doc.add_heading(line[4:].strip(), level=3)
            set_heading_style(p, 3)
            i += 1
            continue
        if line.startswith("## "):
            p = doc.add_heading(line[3:].strip(), level=2)
            set_heading_style(p, 2)
            i += 1
            continue
        if line.startswith("# "):
            p = doc.add_heading(line[2:].strip(), level=1)
            set_heading_style(p, 1)
            i += 1
            continue

        # Horizontal rule
        if line.strip() in ("---", "***", "___"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "AAAAAA")
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # Blockquote
        if line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            inline_bold(p, line[2:].strip())
            for run in p.runs:
                run.font.italic = True
                run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            i += 1
            continue

        # Italic line (e.g. *All numbers cited...*)
        if re.match(r"^\*[^*].+[^*]\*\s*$", line.strip()):
            p = doc.add_paragraph()
            run = p.add_run(line.strip().strip("*"))
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
            run.font.size = Pt(9)
            i += 1
            continue

        # Bullet / list
        if re.match(r"^[-*] ", line):
            p = doc.add_paragraph(style="List Bullet")
            inline_bold(p, line[2:].strip())
            i += 1
            continue
        if re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            inline_bold(p, re.sub(r"^\d+\. ", "", line).strip())
            i += 1
            continue

        # Empty line
        if line.strip() == "":
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        inline_bold(p, line.strip())
        i += 1

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
